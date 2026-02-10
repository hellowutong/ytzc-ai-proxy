"""
知识库管理器 - 文档管理、向量化、检索
"""

import logging
import os
import shutil
from typing import Optional, List, Dict, Any, Tuple
from datetime import datetime
from dataclasses import dataclass, field
import uuid
from pathlib import Path


logger = logging.getLogger(__name__)


@dataclass
class DocumentChunk:
    """文档片段"""
    index: int
    content: str
    vector_id: Optional[str] = None
    vectorized: bool = False


@dataclass
class Document:
    """文档模型"""
    id: str
    filename: str
    file_type: str  # pdf/txt/doc/jpg/md
    source: str  # upload/rss/conversation
    virtual_model: Optional[str]
    is_shared: bool
    vectorized: bool
    chunk_count: int
    file_path: str
    file_size: int
    upload_time: datetime
    chunks: List[DocumentChunk] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)


@dataclass
class SearchResult:
    """搜索结果"""
    document_id: str
    chunk_index: int
    content: str
    score: float
    source: str
    filename: str


class KnowledgeManager:
    """知识库管理器"""
    
    def __init__(
        self, 
        mongodb_client, 
        qdrant_client, 
        embedding_service=None,
        skill_manager=None,
        config_manager=None
    ):
        """
        初始化知识库管理器
        
        Args:
            mongodb_client: MongoDB客户端实例
            qdrant_client: Qdrant客户端实例
            embedding_service: Embedding服务实例（可选）
            skill_manager: Skill管理器实例（可选）
            config_manager: 配置管理器实例（可选）
        """
        self._mongodb = mongodb_client
        self._qdrant = qdrant_client
        self._embedding_service = embedding_service
        self._skill_manager = skill_manager
        self._config_manager = config_manager
        
        # 获取数据库和集合
        if config_manager:
            db_name = config_manager.get('storage.mongodb.database', 'ai_gateway')
        else:
            db_name = 'ai_gateway'
        
        self._db = self._mongodb[db_name]
        self._docs_collection = self._db["knowledge_docs"]
        
        # 向量集合名称
        self._collection_name = "knowledge_base"
    
    def _generate_uuid(self) -> str:
        """生成UUID"""
        return str(uuid.uuid4())
    
    async def upload_document(
        self,
        file_content: bytes,
        filename: str,
        virtual_model: str,
        is_shared: bool = False,
        chunk_size: int = 500,
        overlap: int = 50,
        language: str = "zh"
    ) -> Document:
        """
        上传文档流程
        
        Args:
            file_content: 文件内容（字节）
            filename: 文件名
            virtual_model: 虚拟模型名称
            is_shared: 是否共享
            chunk_size: 分块大小
            overlap: 重叠大小
            language: 语言
            
        Returns:
            Document: 文档对象
        """
        try:
            # 1. 生成文档ID和保存路径
            document_id = self._generate_uuid()
            file_ext = Path(filename).suffix.lower()
            
            # 确定文件类型
            file_type = self._get_file_type(file_ext)
            
            # 构建保存路径
            upload_dir = Path("../upload/textfile")
            upload_dir.mkdir(parents=True, exist_ok=True)
            file_path = upload_dir / f"{document_id}{file_ext}"
            
            # 2. 保存文件
            with open(file_path, "wb") as f:
                f.write(file_content)
            
            # 3. 提取文本内容
            text_content = await self._extract_text(file_path, file_type)
            
            # 4. 分段处理
            chunks = self._split_text(text_content, chunk_size, overlap)
            chunk_objects = [
                DocumentChunk(index=i, content=chunk)
                for i, chunk in enumerate(chunks)
            ]
            
            # 5. 保存元数据到MongoDB
            now = datetime.utcnow()
            doc_data = {
                "_id": document_id,
                "filename": filename,
                "type": file_type,
                "source": "upload",
                "virtual_model": virtual_model,
                "is_shared": is_shared,
                "vectorized": False,
                "chunk_count": len(chunks),
                "file_path": str(file_path),
                "file_size": len(file_content),
                "upload_time": now,
                "chunks": [
                    {"index": c.index, "content": c.content, "vector_id": None, "vectorized": False}
                    for c in chunk_objects
                ],
                "metadata": {
                    "chunk_size": chunk_size,
                    "overlap": overlap,
                    "language": language
                }
            }
            
            await self._docs_collection.insert_one(doc_data)
            
            # 6. 向量化并存储到Qdrant
            if self._embedding_service:
                await self.vectorize_document(document_id)
            
            logger.info(f"文档上传成功: {document_id}, 文件名: {filename}")
            
            return Document(
                id=document_id,
                filename=filename,
                file_type=file_type,
                source="upload",
                virtual_model=virtual_model,
                is_shared=is_shared,
                vectorized=False,
                chunk_count=len(chunks),
                file_path=str(file_path),
                file_size=len(file_content),
                upload_time=now,
                chunks=chunk_objects
            )
            
        except Exception as e:
            logger.error(f"文档上传失败: {e}")
            raise
    
    def _get_file_type(self, extension: str) -> str:
        """根据扩展名获取文件类型"""
        type_map = {
            ".pdf": "pdf",
            ".txt": "txt",
            ".doc": "doc",
            ".docx": "doc",
            ".md": "md",
            ".jpg": "jpg",
            ".jpeg": "jpg",
            ".png": "png"
        }
        return type_map.get(extension, "unknown")
    
    async def _extract_text(self, file_path: Path, file_type: str) -> str:
        """从文件中提取文本"""
        try:
            if file_type == "txt" or file_type == "md":
                with open(file_path, "r", encoding="utf-8") as f:
                    return f.read()
            
            elif file_type == "pdf":
                # 使用PyPDF2或其他库提取PDF文本
                try:
                    import PyPDF2
                    text = ""
                    with open(file_path, "rb") as f:
                        reader = PyPDF2.PdfReader(f)
                        for page in reader.pages:
                            text += page.extract_text() + "\n"
                    return text
                except ImportError:
                    logger.warning("PyPDF2未安装，无法提取PDF文本")
                    return ""
            
            elif file_type == "doc":
                # 使用python-docx提取Word文本
                try:
                    import docx
                    doc = docx.Document(file_path)
                    return "\n".join([para.text for para in doc.paragraphs])
                except ImportError:
                    logger.warning("python-docx未安装，无法提取Word文本")
                    return ""
            
            elif file_type in ["jpg", "png"]:
                # OCR图像提取（需要tesseract）
                logger.warning("OCR功能需要额外配置")
                return ""
            
            else:
                return ""
                
        except Exception as e:
            logger.error(f"文本提取失败: {e}")
            return ""
    
    def _split_text(self, text: str, chunk_size: int, overlap: int) -> List[str]:
        """将文本分块"""
        if not text:
            return []
        
        chunks = []
        start = 0
        text_length = len(text)
        
        while start < text_length:
            end = min(start + chunk_size, text_length)
            
            # 尝试在句子或段落边界处分割
            if end < text_length:
                # 查找最近的句子结束符
                for sep in ["\n\n", ".", "!", "?", "；", "。", "！", "？"]:
                    pos = text.rfind(sep, start, end)
                    if pos > start:
                        end = pos + len(sep)
                        break
            
            chunks.append(text[start:end].strip())
            start = end - overlap if end - overlap > start else end
        
        return chunks
    
    async def vectorize_document(self, document_id: str) -> bool:
        """
        重新向量化文档
        
        Args:
            document_id: 文档ID
            
        Returns:
            bool: 是否成功
        """
        try:
            # 获取文档信息
            doc = await self._docs_collection.find_one({"_id": document_id})
            if not doc:
                logger.warning(f"文档不存在: {document_id}")
                return False
            
            if not self._embedding_service:
                logger.warning("Embedding服务未配置")
                return False
            
            chunks = doc.get("chunks", [])
            if not chunks:
                logger.warning(f"文档没有内容块: {document_id}")
                return False
            
            # 确保Qdrant集合存在
            await self._ensure_collection()
            
            # 为每个块生成向量并存储
            vector_points = []
            updated_chunks = []
            
            for chunk in chunks:
                chunk_text = chunk.get("content", "")
                if not chunk_text:
                    continue
                
                # 生成embedding
                vector = await self._embedding_service.embed(chunk_text)
                vector_id = chunk.get("vector_id") or self._generate_uuid()
                
                vector_points.append({
                    "id": vector_id,
                    "vector": vector,
                    "payload": {
                        "document_id": document_id,
                        "chunk_index": chunk.get("index", 0),
                        "virtual_model": doc.get("virtual_model"),
                        "is_shared": doc.get("is_shared", False),
                        "source": doc.get("source", "upload"),
                        "created_at": doc.get("upload_time", datetime.utcnow()).isoformat(),
                        "text_preview": chunk_text[:100]
                    }
                })
                
                updated_chunks.append({
                    "index": chunk.get("index", 0),
                    "content": chunk_text,
                    "vector_id": vector_id,
                    "vectorized": True
                })
            
            # 批量插入到Qdrant
            if vector_points:
                self._qdrant.upsert(
                    collection_name=self._collection_name,
                    points=vector_points
                )
            
            # 更新MongoDB
            await self._docs_collection.update_one(
                {"_id": document_id},
                {
                    "$set": {
                        "vectorized": True,
                        "chunks": updated_chunks,
                        "updated_at": datetime.utcnow()
                    }
                }
            )
            
            logger.info(f"文档向量化成功: {document_id}, 块数: {len(vector_points)}")
            return True
            
        except Exception as e:
            logger.error(f"文档向量化失败: {e}")
            return False
    
    async def _ensure_collection(self):
        """确保Qdrant集合存在"""
        try:
            collections = self._qdrant.get_collections().collections
            collection_names = [c.name for c in collections]
            
            if self._collection_name not in collection_names:
                # 获取向量维度配置
                dimension = 1024  # 默认BAAI/bge-m3维度
                if self._config_manager:
                    dimension = self._config_manager.get(
                        "ai-gateway.knowledge.vector_dimension", 1024
                    )
                
                self._qdrant.create_collection(
                    collection_name=self._collection_name,
                    vectors_config={
                        "size": dimension,
                        "distance": "Cosine"
                    }
                )
                
                logger.info(f"创建Qdrant集合: {self._collection_name}")
                
        except Exception as e:
            logger.error(f"确保集合存在失败: {e}")
    
    async def search(
        self,
        query: str,
        virtual_model: Optional[str] = None,
        threshold: float = 0.76,
        top_k: int = 5
    ) -> List[SearchResult]:
        """
        向量检索
        
        Args:
            query: 查询文本
            virtual_model: 虚拟模型名称筛选
            threshold: 相似度阈值
            top_k: 返回结果数量
            
        Returns:
            List[SearchResult]: 搜索结果列表
        """
        try:
            if not self._embedding_service:
                logger.warning("Embedding服务未配置")
                return []
            
            # 1. 获取query的embedding
            query_vector = await self._embedding_service.embed(query)
            
            # 2. 构建过滤条件
            filters = None
            if virtual_model:
                filters = {
                    "should": [
                        {"key": "virtual_model", "match": {"value": virtual_model}},
                        {"key": "is_shared", "match": {"value": True}}
                    ]
                }
            
            # 3. 在Qdrant中搜索
            results = self._qdrant.search(
                collection_name=self._collection_name,
                query_vector=query_vector,
                limit=top_k * 2,  # 多取一些，过滤后再限制
                query_filter=filters
            )
            
            # 4. 过滤threshold并构建结果
            search_results = []
            for result in results:
                score = result.score
                if score >= threshold:
                    payload = result.payload
                    search_results.append(SearchResult(
                        document_id=payload.get("document_id", ""),
                        chunk_index=payload.get("chunk_index", 0),
                        content=payload.get("text_preview", ""),
                        score=score,
                        source=payload.get("source", ""),
                        filename=payload.get("filename", "")
                    ))
                
                if len(search_results) >= top_k:
                    break
            
            return search_results
            
        except Exception as e:
            logger.error(f"向量检索失败: {e}")
            return []
    
    async def extract_knowledge(
        self, 
        text: str, 
        virtual_model: str
    ) -> List[Dict[str, Any]]:
        """
        执行知识提取Skill
        
        Args:
            text: 待提取文本
            virtual_model: 虚拟模型名称
            
        Returns:
            List[Dict[str, Any]]: 知识块列表
        """
        try:
            if not self._skill_manager:
                logger.warning("Skill管理器未配置")
                return []
            
            # 调用Skill执行知识提取
            skill_result = await self._skill_manager.execute(
                "knowledge", "简单提取",
                text=text,
                virtual_model=virtual_model
            )
            
            return skill_result.get("chunks", [])
            
        except Exception as e:
            logger.error(f"知识提取失败: {e}")
            return []
    
    async def list_documents(
        self,
        virtual_model: Optional[str] = None,
        source: Optional[str] = None,
        vectorized: Optional[bool] = None,
        limit: int = 20,
        offset: int = 0
    ) -> Tuple[List[Document], int]:
        """
        列出文档
        
        Args:
            virtual_model: 虚拟模型筛选
            source: 来源筛选
            vectorized: 是否已向量化筛选
            limit: 返回数量限制
            offset: 偏移量
            
        Returns:
            Tuple[List[Document], int]: (文档列表, 总数)
        """
        try:
            # 构建查询条件
            query = {}
            if virtual_model:
                query["virtual_model"] = virtual_model
            if source:
                query["source"] = source
            if vectorized is not None:
                query["vectorized"] = vectorized
            
            # 查询数据
            cursor = self._docs_collection.find(query) \
                .sort("upload_time", -1) \
                .skip(offset) \
                .limit(limit)
            
            documents = []
            async for doc in cursor:
                chunks = [
                    DocumentChunk(
                        index=c.get("index", 0),
                        content=c.get("content", ""),
                        vector_id=c.get("vector_id"),
                        vectorized=c.get("vectorized", False)
                    )
                    for c in doc.get("chunks", [])
                ]
                
                documents.append(Document(
                    id=doc["_id"],
                    filename=doc.get("filename", ""),
                    file_type=doc.get("type", ""),
                    source=doc.get("source", ""),
                    virtual_model=doc.get("virtual_model"),
                    is_shared=doc.get("is_shared", False),
                    vectorized=doc.get("vectorized", False),
                    chunk_count=doc.get("chunk_count", 0),
                    file_path=doc.get("file_path", ""),
                    file_size=doc.get("file_size", 0),
                    upload_time=doc.get("upload_time", datetime.utcnow()),
                    chunks=chunks,
                    metadata=doc.get("metadata", {})
                ))
            
            total = await self._docs_collection.count_documents(query)
            
            return documents, total
            
        except Exception as e:
            logger.error(f"列出文档失败: {e}")
            return [], 0
    
    async def delete_document(self, document_id: str) -> bool:
        """
        删除文档
        
        Args:
            document_id: 文档ID
            
        Returns:
            bool: 是否成功
        """
        try:
            # 获取文档信息
            doc = await self._docs_collection.find_one({"_id": document_id})
            if not doc:
                logger.warning(f"文档不存在: {document_id}")
                return False
            
            # 删除Qdrant中的向量
            vector_ids = [
                chunk.get("vector_id") 
                for chunk in doc.get("chunks", [])
                if chunk.get("vector_id")
            ]
            
            if vector_ids:
                self._qdrant.delete(
                    collection_name=self._collection_name,
                    points_selector=vector_ids
                )
            
            # 删除文件
            file_path = doc.get("file_path")
            if file_path and os.path.exists(file_path):
                os.remove(file_path)
            
            # 删除MongoDB记录
            result = await self._docs_collection.delete_one({"_id": document_id})
            
            success = result.deleted_count > 0
            if success:
                logger.info(f"删除文档成功: {document_id}")
            
            return success
            
        except Exception as e:
            logger.error(f"删除文档失败: {e}")
            return False
    
    async def get_document(self, document_id: str) -> Optional[Document]:
        """
        获取文档详情
        
        Args:
            document_id: 文档ID
            
        Returns:
            Optional[Document]: 文档对象，不存在时返回None
        """
        try:
            doc = await self._docs_collection.find_one({"_id": document_id})
            
            if not doc:
                return None
            
            chunks = [
                DocumentChunk(
                    index=c.get("index", 0),
                    content=c.get("content", ""),
                    vector_id=c.get("vector_id"),
                    vectorized=c.get("vectorized", False)
                )
                for c in doc.get("chunks", [])
            ]
            
            return Document(
                id=doc["_id"],
                filename=doc.get("filename", ""),
                file_type=doc.get("type", ""),
                source=doc.get("source", ""),
                virtual_model=doc.get("virtual_model"),
                is_shared=doc.get("is_shared", False),
                vectorized=doc.get("vectorized", False),
                chunk_count=doc.get("chunk_count", 0),
                file_path=doc.get("file_path", ""),
                file_size=doc.get("file_size", 0),
                upload_time=doc.get("upload_time", datetime.utcnow()),
                chunks=chunks,
                metadata=doc.get("metadata", {})
            )
            
        except Exception as e:
            logger.error(f"获取文档失败: {e}")
            return None

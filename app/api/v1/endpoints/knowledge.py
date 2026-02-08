"""
Knowledge base API endpoints for document management and extraction.

Features:
- Document CRUD operations
- File upload/download
- Text extraction
- Knowledge chunk management
- Vector search
"""

import asyncio
import math
import os
import random
import re  # Add re import for regex pattern matching
import shutil
import time
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

from fastapi import APIRouter, Depends, File, Form, HTTPException, Query, UploadFile, status
from fastapi.responses import FileResponse

from core.config_manager import get_config
from core.logger import get_logger
from db.mongodb import get_mongodb_client
from db.qdrant import get_qdrant_client
from models.knowledge import (
    Document,
    DocumentCreateRequest,
    DocumentListRequest,
    DocumentListResponse,
    DocumentStats,
    DocumentStatus,
    DocumentType,
    DocumentUpdateRequest,
    ExtractionRequest,
    ExtractionResult,
    FileUploadResponse,
    KnowledgeChunk,
    KnowledgeQueryRequest,
    KnowledgeQueryResponse,
)

router = APIRouter(prefix="/knowledge", tags=["Knowledge Base"])
logger = get_logger("knowledge_api")

# Upload directory from config
UPLOAD_DIR = Path("./uploads/knowledge")
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

# Content type mapping
CONTENT_TYPE_MAP = {
    # Text
    ".txt": DocumentType.TEXT,
    ".md": DocumentType.TEXT,
    ".markdown": DocumentType.TEXT,
    ".json": DocumentType.TEXT,
    ".yaml": DocumentType.TEXT,
    ".yml": DocumentType.TEXT,
    # PDF
    ".pdf": DocumentType.PDF,
    # Documents
    ".doc": DocumentType.DOC,
    ".docx": DocumentType.DOC,
    # Images
    ".jpg": DocumentType.IMAGE,
    ".jpeg": DocumentType.IMAGE,
    ".png": DocumentType.IMAGE,
    ".gif": DocumentType.IMAGE,
    ".webp": DocumentType.IMAGE,
    # Video
    ".mp4": DocumentType.VIDEO,
    ".avi": DocumentType.VIDEO,
    ".mov": DocumentType.VIDEO,
    ".wmv": DocumentType.VIDEO,
    ".mkv": DocumentType.VIDEO,
    # Audio
    ".mp3": DocumentType.AUDIO,
    ".wav": DocumentType.AUDIO,
    ".ogg": DocumentType.AUDIO,
    ".m4a": DocumentType.AUDIO,
    ".aac": DocumentType.AUDIO,
}


def get_content_type(filename: str) -> DocumentType:
    """Determine content type from file extension."""
    ext = Path(filename).suffix.lower()
    return CONTENT_TYPE_MAP.get(ext, DocumentType.TEXT)


async def extract_text_from_pdf(file_path: str) -> tuple[str, float]:
    """
    Extract text from PDF file using pypdf.
    Returns (text, confidence_score).
    """
    try:
        from pypdf import PdfReader
        
        reader = PdfReader(file_path)
        text_parts = []
        
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        
        text = "\n".join(text_parts)
        
        # Calculate confidence based on extracted text
        if not text:
            return "", 0.0
        
        # Confidence based on text density and completeness
        total_chars = len(text)
        avg_chars_per_page = total_chars / len(reader.pages) if reader.pages else 0
        
        # High confidence if reasonable amount of text per page (>100 chars)
        if avg_chars_per_page > 100:
            confidence = 0.95
        elif avg_chars_per_page > 50:
            confidence = 0.85
        elif avg_chars_per_page > 10:
            confidence = 0.70
        else:
            confidence = 0.50
        
        logger.info(f"Extracted {total_chars} chars from {len(reader.pages)} pages, confidence: {confidence}")
        return text, confidence
        
    except Exception as e:
        logger.error(f"PDF extraction failed: {e}")
        return f"[PDF extraction error: {str(e)}]", 0.0


async def extract_text_from_docx(file_path: str) -> tuple[str, float]:
    """
    Extract text from DOCX file using python-docx.
    Returns (text, confidence_score).
    """
    try:
        from docx import Document as DocxDocument
        
        doc = DocxDocument(file_path)
        text_parts = []
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Extract table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        text_parts.append(cell_text)
        
        text = "\n".join(text_parts)
        
        # Calculate confidence
        if not text:
            return "", 0.0
        
        paragraph_count = len([p for p in doc.paragraphs if p.text.strip()])
        table_count = len(doc.tables)
        
        # High confidence for documents with substantial content
        if len(text) > 1000:
            confidence = 0.95
        elif len(text) > 100:
            confidence = 0.90
        elif len(text) > 10:
            confidence = 0.75
        else:
            confidence = 0.50
        
        logger.info(f"Extracted {len(text)} chars from {paragraph_count} paragraphs, {table_count} tables, confidence: {confidence}")
        return text, confidence
        
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        return f"[DOCX extraction error: {str(e)}]", 0.0


async def save_upload_file(upload_file: UploadFile, document_id: str) -> tuple[str, int]:
    """Save uploaded file and return path and size."""
    file_ext = Path(upload_file.filename or "unknown").suffix
    file_path = UPLOAD_DIR / f"{document_id}{file_ext}"
    
    content = await upload_file.read()
    file_size = len(content)
    
    with open(file_path, "wb") as f:
        f.write(content)
    
    return str(file_path), file_size


async def extract_text_from_file(file_path: str, content_type: DocumentType) -> tuple[str, float]:
    """
    Extract text from file based on content type.
    Returns (text, confidence_score).
    """
    # For text files, just read the content
    if content_type == DocumentType.TEXT:
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
            return content, 1.0
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, "r", encoding="latin-1") as f:
                content = f.read()
            return content, 0.9
    
    # For other types, this would integrate with specialized extractors
    # For now, return placeholder indicating extraction needed
    if content_type == DocumentType.PDF:
        # Extract text using pypdf
        if os.path.exists(file_path):
            return await extract_text_from_pdf(file_path)
        return f"[PDF file not found: {file_path}]", 0.0
    
    if content_type == DocumentType.DOC:
        # Extract text using python-docx
        if os.path.exists(file_path):
            return await extract_text_from_docx(file_path)
        return f"[DOCX file not found: {file_path}]", 0.0
    
    if content_type in [DocumentType.IMAGE, DocumentType.VIDEO, DocumentType.AUDIO]:
        # These would use OCR/transcription
        return f"[{content_type.value} content at {file_path} - transcription pending]", 0.0
    
    return f"[Unsupported content type: {content_type.value}]", 0.0


async def create_knowledge_chunks(
    document_id: str,
    content: str,
    model_id: Optional[str],
    topics: List[str]
) -> List[KnowledgeChunk]:
    """
    Create knowledge chunks from document content.
    Simple implementation: split by paragraphs/sections.
    """
    chunks = []
    
    # Simple chunking strategy: split by double newlines
    # In production, this would use more sophisticated chunking
    raw_chunks = [c.strip() for c in content.split("\n\n") if c.strip()]
    
    # If chunks are too large, split further
    max_chunk_size = 1000  # characters
    final_chunks = []
    
    for chunk in raw_chunks:
        if len(chunk) <= max_chunk_size:
            final_chunks.append(chunk)
        else:
            # Split large chunks by sentences or fixed size
            for i in range(0, len(chunk), max_chunk_size):
                final_chunks.append(chunk[i:i+max_chunk_size])
    
    # Create chunk objects
    for idx, chunk_text in enumerate(final_chunks):
        chunk = KnowledgeChunk(
            id=f"{document_id}_chunk_{idx}",
            content=chunk_text,
            document_id=document_id,
            model_id=model_id,
            topic=topics[0] if topics else None,
            metadata={
                "chunk_index": idx,
                "total_chunks": len(final_chunks),
                "char_count": len(chunk_text),
            }
        )
        chunks.append(chunk)
    
    return chunks


async def generate_embeddings(chunks: List[KnowledgeChunk]) -> List[List[float]]:
    """
    Generate embeddings for knowledge chunks using embedding API.
    Supports OpenAI-compatible APIs including SiliconFlow.
    """
    if not chunks:
        return []
    
    # Get embedding config
    embedding_config = get_config("ai-gateway.knowledge.embedding", {})
    model = embedding_config.get("model", "BAAI/bge-m3")
    base_url = embedding_config.get("base_url", "https://api.siliconflow.cn/v1")
    api_key = embedding_config.get("api_key", "")
    
    # Fallback to random embeddings if no API key configured
    if not api_key:
        logger.warning("No embedding API key configured, using random embeddings")
        import random
        import math
        embeddings = []
        for _ in chunks:
            embedding = [random.uniform(-1, 1) for _ in range(768)]
            magnitude = math.sqrt(sum(x**2 for x in embedding))
            embedding = [x / magnitude for x in embedding]
            embeddings.append(embedding)
        return embeddings
    
    # Extract text content from chunks
    texts = [chunk.content for chunk in chunks]
    
    try:
        from openai import AsyncOpenAI
        import math
        
        # Create async client
        client = AsyncOpenAI(base_url=base_url, api_key=api_key)
        
        # Batch processing (API limit: 100 texts per request)
        batch_size = 100
        all_embeddings = []
        
        for i in range(0, len(texts), batch_size):
            batch_texts = texts[i:i + batch_size]
            
            # Retry logic with exponential backoff
            max_retries = 3
            retry_delay = 1.0
            
            for attempt in range(max_retries):
                try:
                    response = await client.embeddings.create(
                        model=model,
                        input=batch_texts
                    )
                    
                    # Extract embeddings from response
                    batch_embeddings = [item.embedding for item in response.data]
                    
                    # Normalize embeddings (ensure unit length)
                    for embedding in batch_embeddings:
                        magnitude = math.sqrt(sum(x**2 for x in embedding))
                        if magnitude > 0:
                            normalized = [x / magnitude for x in embedding]
                            all_embeddings.append(normalized)
                        else:
                            all_embeddings.append(embedding)
                    
                    logger.info(f"Generated embeddings for batch {i//batch_size + 1}, {len(batch_texts)} texts")
                    break
                    
                except Exception as e:
                    if attempt < max_retries - 1:
                        logger.warning(f"Embedding API retry {attempt + 1}/{max_retries}: {e}")
                        await asyncio.sleep(retry_delay)
                        retry_delay *= 2
                    else:
                        # Fallback to random embeddings on failure
                        logger.error(f"Embedding API failed after {max_retries} attempts, using fallback")
                        for _ in batch_texts:
                            embedding = [random.uniform(-1, 1) for _ in range(768)]
                            magnitude = math.sqrt(sum(x**2 for x in embedding))
                            normalized = [x / magnitude for x in embedding]
                            all_embeddings.append(normalized)
        
        return all_embeddings
        
    except ImportError:
        logger.warning("OpenAI client not installed, using random embeddings")
        import random
        import math
        embeddings = []
        for _ in chunks:
            embedding = [random.uniform(-1, 1) for _ in range(768)]
            magnitude = math.sqrt(sum(x**2 for x in embedding))
            embedding = [x / magnitude for x in embedding]
            embeddings.append(embedding)
        return embeddings
        
    except Exception as e:
        logger.error(f"Embedding generation failed: {e}, using random fallback")
        embeddings = []
        for _ in chunks:
            embedding = [random.uniform(-1, 1) for _ in range(768)]
            magnitude = math.sqrt(sum(x**2 for x in embedding))
            embedding = [x / magnitude for x in embedding]
            embeddings.append(embedding)
        return embeddings


async def generate_query_embedding(query_text: str) -> List[float]:
    """
    Generate embedding for a query string using embedding API.
    """
    if not query_text:
        return [0.0] * 768
    
    # Get embedding config
    embedding_config = get_config("ai-gateway.knowledge.embedding", {})
    model = embedding_config.get("model", "BAAI/bge-m3")
    base_url = embedding_config.get("base_url", "https://api.siliconflow.cn/v1")
    api_key = embedding_config.get("api_key", "")
    
    # Fallback to random if no API key
    if not api_key:
        logger.warning("No embedding API key configured, using random query embedding")
        embedding = [random.uniform(-1, 1) for _ in range(768)]
        magnitude = math.sqrt(sum(x**2 for x in embedding))
        return [x / magnitude for x in embedding] if magnitude > 0 else embedding
    
    try:
        from openai import AsyncOpenAI
        
        client = AsyncOpenAI(base_url=base_url, api_key=api_key)
        
        # Retry logic
        max_retries = 3
        retry_delay = 1.0
        
        for attempt in range(max_retries):
            try:
                response = await client.embeddings.create(
                    model=model,
                    input=[query_text]
                )
                
                embedding = response.data[0].embedding
                
                # Normalize
                magnitude = math.sqrt(sum(x**2 for x in embedding))
                if magnitude > 0:
                    return [x / magnitude for x in embedding]
                return embedding
                
            except Exception as e:
                if attempt < max_retries - 1:
                    logger.warning(f"Query embedding retry {attempt + 1}/{max_retries}: {e}")
                    await asyncio.sleep(retry_delay)
                    retry_delay *= 2
                else:
                    logger.error(f"Query embedding failed after {max_retries} attempts")
        
        # Fallback
        embedding = [random.uniform(-1, 1) for _ in range(768)]
        magnitude = math.sqrt(sum(x**2 for x in embedding))
        return [x / magnitude for x in embedding] if magnitude > 0 else embedding
        
    except ImportError:
        logger.warning("OpenAI client not installed, using random query embedding")
        embedding = [random.uniform(-1, 1) for _ in range(768)]
        magnitude = math.sqrt(sum(x**2 for x in embedding))
        return [x / magnitude for x in embedding] if magnitude > 0 else embedding
        
    except Exception as e:
        logger.error(f"Query embedding error: {e}, using random fallback")
        embedding = [random.uniform(-1, 1) for _ in range(768)]
        magnitude = math.sqrt(sum(x**2 for x in embedding))
        return [x / magnitude for x in embedding] if magnitude > 0 else embedding


async def classify_topic(content: str) -> str:
    """
    Classify content into a topic based on keyword matching.
    Uses patterns defined in config (ai-gateway.knowledge.topics).
    """
    if not content or not content.strip():
        return "general"
    
    # Get topics configuration
    topics_config = get_config("ai-gateway.knowledge.topics", {})
    
    # Collect scores for each topic
    topic_scores = {}
    
    # Process auto_classify patterns
    for topic_rule in topics_config.get("auto_classify", []):
        topic_name = topic_rule.get("topic", "general")
        patterns = topic_rule.get("patterns", [])
        
        score = 0
        for pattern in patterns:
            # Count occurrences of pattern (case-insensitive)
            count = len(re.findall(pattern, content, re.IGNORECASE))
            score += count
        
        if score > 0:
            topic_scores[topic_name] = topic_scores.get(topic_name, 0) + score
    
    # Process self_classify patterns
    for topic_rule in topics_config.get("self_classify", []):
        topic_name = topic_rule.get("topic", "general")
        patterns = topic_rule.get("patterns", [])
        
        score = 0
        for pattern in patterns:
            count = len(re.findall(pattern, content, re.IGNORECASE))
            score += count
        
        if score > 0:
            topic_scores[topic_name] = topic_scores.get(topic_name, 0) + score
    
    # Return topic with highest score, or "general" if no match
    if topic_scores:
        best_topic = max(topic_scores, key=topic_scores.get)
        logger.info(f"Classified topic: {best_topic} (score: {topic_scores[best_topic]})")
        return best_topic
    
    return "general"


# ============================================================================
# API Endpoints
# ============================================================================

@router.get("/documents", response_model=DocumentListResponse)
async def list_documents(
    model_id: Optional[str] = Query(None, description="Filter by model ID"),
    status: Optional[DocumentStatus] = Query(None, description="Filter by status"),
    content_type: Optional[DocumentType] = Query(None, description="Filter by type"),
    shared: Optional[bool] = Query(None, description="Filter by shared status"),
    search: Optional[str] = Query(None, description="Search in title/content"),
    skip: int = Query(0, ge=0),
    limit: int = Query(20, ge=1, le=100),
):
    """List documents with filters and pagination."""
    try:
        db = get_mongodb_client().db
        
        # Build query
        query = {}
        if model_id:
            query["$or"] = [{"model_id": model_id}, {"shared": True}]
        if status:
            query["status"] = status.value
        if content_type:
            query["content_type"] = content_type.value
        if shared is not None:
            query["shared"] = shared
        if search:
            query["$or"] = [
                {"title": {"$regex": search, "$options": "i"}},
                {"content": {"$regex": search, "$options": "i"}},
            ]
        
        # Get total count
        total = await db.documents.count_documents(query)
        
        # Get documents
        cursor = db.documents.find(query).skip(skip).limit(limit).sort("created_at", -1)
        documents = []
        async for doc in cursor:
            doc["_id"] = str(doc["_id"])
            documents.append(Document(**doc))
        
        return DocumentListResponse(
            total=total,
            documents=documents,
            skip=skip,
            limit=limit,
        )
    
    except Exception as e:
        logger.error(f"Failed to list documents: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to list documents: {str(e)}")


@router.post("/documents", response_model=Document)
async def create_document(request: DocumentCreateRequest):
    """Create a new document record (without file upload)."""
    try:
        db = get_mongodb_client().db
        
        document = Document(
            title=request.title,
            content_type=request.content_type,
            model_id=request.model_id,
            shared=request.shared,
            tags=request.tags,
            source_url=request.source_url,
            source_type=request.source_type or "manual",
            custom_metadata=request.custom_metadata,
            status=DocumentStatus.PENDING,
        )
        
        result = await db.documents.insert_one(document.model_dump(exclude={"id"}, by_alias=True))
        document.id = str(result.inserted_id)
        
        logger.info(f"Created document: {document.id}")
        return document
    
    except Exception as e:
        logger.error(f"Failed to create document: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to create document: {str(e)}")


@router.get("/documents/{document_id}", response_model=Document)
async def get_document(document_id: str):
    """Get a specific document by ID."""
    try:
        from bson.objectid import ObjectId
        db = get_mongodb_client().db
        
        doc = await db.documents.find_one({"_id": ObjectId(document_id)})
        if not doc:
            raise HTTPException(status_code=404, detail="Document not found")
        
        doc["_id"] = str(doc["_id"])
        return Document(**doc)
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to get document: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to get document: {str(e)}")


@router.put("/documents/{document_id}", response_model=Document)
async def update_document(document_id: str, request: DocumentUpdateRequest):
    """Update document metadata."""
    try:
        from bson.objectid import ObjectId
        db = get_mongodb_client().db
        
        # Check if document exists
        existing = await db.documents.find_one({"_id": ObjectId(document_id)})
        if not existing:
            raise HTTPException(status_code=404, detail="Document not found")
        
        # Build update data (exclude None values)
        update_data = {}
        for field, value in request.model_dump(exclude_none=True).items():
            update_data[field] = value
        
        update_data["updated_at"] = datetime.utcnow()
        
        await db.documents.update_one(
            {"_id": ObjectId(document_id)},
            {"$set": update_data}
        )
        
        # Return updated document
        doc = await db.documents.find_one({"_id": ObjectId(document_id)})
        doc["_id"] = str(doc["_id"])
        
        logger.info(f"Updated document: {document_id}")
        return Document(**doc)
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to update document: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to update document: {str(e)}")


@router.delete("/documents/{document_id}")
async def delete_document(document_id: str):
    """Delete a document and its associated data."""
    try:
        from bson.objectid import ObjectId
        db = get_mongodb_client().db
        
        # Get document to find file path
        doc = await db.documents.find_one({"_id": ObjectId(document_id)})
        if not doc:
            raise HTTPException(status_code=404, detail="Document not found")
        
        # Delete associated file if exists
        if doc.get("file_path") and os.path.exists(doc["file_path"]):
            os.remove(doc["file_path"])
        
        # Delete from Qdrant (knowledge chunks)
        try:
            qdrant = get_qdrant_client()
            collection = qdrant.get_collection()
            # Delete all chunks for this document
            qdrant.client.delete(
                collection_name=collection,
                points_selector={"filter": {"must": [{"key": "document_id", "match": {"value": document_id}}]}}
            )
        except Exception as e:
            logger.warning(f"Failed to delete from Qdrant: {e}")
        
        # Delete from MongoDB
        await db.documents.delete_one({"_id": ObjectId(document_id)})
        await db.knowledge_chunks.delete_many({"document_id": document_id})
        
        logger.info(f"Deleted document: {document_id}")
        return {"success": True, "message": "Document deleted successfully"}
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to delete document: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to delete document: {str(e)}")


@router.post("/documents/upload", response_model=FileUploadResponse)
async def upload_file(
    file: UploadFile = File(..., description="File to upload"),
    model_id: Optional[str] = Form(None, description="Associate with model"),
    shared: bool = Form(False, description="Share across models"),
    title: Optional[str] = Form(None, description="Custom title (defaults to filename)"),
    tags: Optional[str] = Form(None, description="Comma-separated tags"),
):
    """Upload a file and create document record."""
    try:
        # Validate file
        if not file.filename:
            raise HTTPException(status_code=400, detail="No filename provided")
        
        # Check file size
        content = await file.read()
        file_size = len(content)
        await file.seek(0)  # Reset for reading again
        
        max_size_mb = get_config("ai-gateway.text.upload.max_size_mb", 100)
        if file_size > max_size_mb * 1024 * 1024:
            raise HTTPException(
                status_code=400, 
                detail=f"File too large. Max size: {max_size_mb}MB"
            )
        
        # Determine content type
        content_type = get_content_type(file.filename)
        
        # Create document record first
        doc_title = title or file.filename
        doc_tags = [t.strip() for t in tags.split(",")] if tags else []
        
        document = Document(
            title=doc_title,
            content_type=content_type,
            file_size=file_size,
            mime_type=file.content_type,
            model_id=model_id,
            shared=shared,
            tags=doc_tags,
            source_type="upload",
            status=DocumentStatus.PENDING,
        )
        
        db = get_mongodb_client().db
        result = await db.documents.insert_one(document.model_dump(exclude={"id"}, by_alias=True))
        document.id = str(result.inserted_id)
        
        # Save file
        file_path, _ = await save_upload_file(file, document.id)
        
        # Update document with file path
        await db.documents.update_one(
            {"_id": result.inserted_id},
            {"$set": {"file_path": file_path}}
        )
        
        logger.info(f"Uploaded file: {file.filename} -> document: {document.id}")
        
        return FileUploadResponse(
            success=True,
            document_id=document.id,
            file_name=file.filename,
            file_size=file_size,
            content_type=content_type,
            message="File uploaded successfully",
            status=DocumentStatus.PENDING,
        )
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to upload file: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to upload file: {str(e)}")


@router.post("/documents/{document_id}/extract", response_model=ExtractionResult)
async def extract_document(document_id: str, request: ExtractionRequest):
    """Extract text from document and create knowledge chunks."""
    start_time = time.time()
    
    try:
        from bson.objectid import ObjectId
        db = get_mongodb_client().db
        
        # Get document
        doc = await db.documents.find_one({"_id": ObjectId(document_id)})
        if not doc:
            raise HTTPException(status_code=404, detail="Document not found")
        
        # Update status
        await db.documents.update_one(
            {"_id": ObjectId(document_id)},
            {"$set": {"status": DocumentStatus.EXTRACTING.value, "updated_at": datetime.utcnow()}}
        )
        
        # Extract text
        file_path = doc.get("file_path")
        content_type = DocumentType(doc.get("content_type", "text"))
        
        if file_path and os.path.exists(file_path):
            content, confidence = await extract_text_from_file(file_path, content_type)
        else:
            content = doc.get("content", "")
            confidence = 1.0 if content else 0.0
        
        # Update with extracted content
        word_count = len(content.split()) if content else 0
        await db.documents.update_one(
            {"_id": ObjectId(document_id)},
            {
                "$set": {
                    "content": content,
                    "status": DocumentStatus.EXTRACTED.value,
                    "extraction_confidence": confidence,
                    "word_count": word_count,
                    "updated_at": datetime.utcnow(),
                }
            }
        )
        
        # Auto-classify topics using keyword matching
        topics = []
        if request.auto_classify:
            topic = await classify_topic(content)
            topics = [topic] if topic != "general" else []
            if not topics:
                topics = ["general"]
        
        # Create knowledge chunks
        await db.documents.update_one(
            {"_id": ObjectId(document_id)},
            {"$set": {"status": DocumentStatus.PROCESSING.value}}
        )
        
        model_id = request.model_id or doc.get("model_id")
        chunks = await create_knowledge_chunks(document_id, content, model_id, topics)
        
        # Generate embeddings
        embeddings = await generate_embeddings(chunks)
        for chunk, embedding in zip(chunks, embeddings):
            chunk.embedding = embedding
        
        # Store in Qdrant
        try:
            qdrant = get_qdrant_client()
            collection = qdrant.get_collection()
            
            points = []
            for chunk in chunks:
                points.append({
                    "id": chunk.id,
                    "vector": chunk.embedding,
                    "payload": {
                        "content": chunk.content,
                        "document_id": chunk.document_id,
                        "model_id": chunk.model_id,
                        "topic": chunk.topic,
                        "metadata": chunk.metadata,
                    }
                })
            
            if points:
                qdrant.client.upsert(collection_name=collection, points=points)
        
        except Exception as e:
            logger.error(f"Failed to store in Qdrant: {e}")
        
        # Store chunks in MongoDB
        if chunks:
            chunk_dicts = [c.model_dump(exclude={"embedding"}) for c in chunks]
            await db.knowledge_chunks.insert_many(chunk_dicts)
        
        # Update document status
        await db.documents.update_one(
            {"_id": ObjectId(document_id)},
            {
                "$set": {
                    "status": DocumentStatus.PROCESSED.value,
                    "chunk_count": len(chunks),
                    "topics": topics,
                    "updated_at": datetime.utcnow(),
                }
            }
        )
        
        processing_time = int((time.time() - start_time) * 1000)
        
        logger.info(f"Extracted document: {document_id}, created {len(chunks)} chunks")
        
        return ExtractionResult(
            document_id=document_id,
            success=True,
            chunks_created=len(chunks),
            topics_detected=topics,
            confidence_score=confidence,
            message=f"Extraction complete. Created {len(chunks)} knowledge chunks.",
            processing_time_ms=processing_time,
        )
    
    except HTTPException:
        # Update status to failed
        try:
            from bson.objectid import ObjectId
            db = get_mongodb_client().db
            await db.documents.update_one(
                {"_id": ObjectId(document_id)},
                {"$set": {"status": DocumentStatus.FAILED.value}}
            )
        except:
            pass
        raise
    except Exception as e:
        logger.error(f"Failed to extract document: {e}")
        # Update status to failed
        try:
            from bson.objectid import ObjectId
            db = get_mongodb_client().db
            await db.documents.update_one(
                {"_id": ObjectId(document_id)},
                {"$set": {"status": DocumentStatus.FAILED.value, "status_message": str(e)}}
            )
        except:
            pass
        raise HTTPException(status_code=500, detail=f"Extraction failed: {str(e)}")


@router.post("/query", response_model=KnowledgeQueryResponse)
async def query_knowledge(request: KnowledgeQueryRequest):
    """Query knowledge base with semantic search."""
    try:
        # Get threshold from config if not provided
        threshold = request.threshold
        if threshold is None:
            threshold = get_config("ai-gateway.knowledge.threshold.retrieval", 0.76)
        
        # Generate embedding for query using actual API
        query_embedding = await generate_query_embedding(request.query)
        
        # Search in Qdrant
        qdrant = get_qdrant_client()
        collection = qdrant.get_collection()
        
        filter_conditions = []
        if request.model_id:
            filter_conditions.append({
                "key": "model_id",
                "match": {"value": request.model_id}
            })
        if request.topic:
            filter_conditions.append({
                "key": "topic",
                "match": {"value": request.topic}
            })
        
        search_result = qdrant.client.search(
            collection_name=collection,
            query_vector=query_embedding,
            query_filter={"must": filter_conditions} if filter_conditions else None,
            limit=request.limit,
            score_threshold=threshold,
        )
        
        # Convert to KnowledgeChunk objects
        chunks = []
        for point in search_result:
            chunk = KnowledgeChunk(
                id=point.id,
                content=point.payload.get("content", ""),
                document_id=point.payload.get("document_id"),
                model_id=point.payload.get("model_id"),
                topic=point.payload.get("topic"),
                metadata=point.payload.get("metadata", {}),
                score=point.score,
            )
            chunks.append(chunk)
        
        return KnowledgeQueryResponse(
            query=request.query,
            results=chunks,
            total_found=len(chunks),
            threshold_used=threshold,
        )
    
    except Exception as e:
        logger.error(f"Failed to query knowledge: {e}")
        raise HTTPException(status_code=500, detail=f"Query failed: {str(e)}")


@router.get("/stats", response_model=DocumentStats)
async def get_stats():
    """Get knowledge base statistics."""
    try:
        db = get_mongodb_client().db
        
        # Total documents
        total = await db.documents.count_documents({})
        
        # By status
        by_status = {}
        for status in DocumentStatus:
            count = await db.documents.count_documents({"status": status.value})
            by_status[status.value] = count
        
        # By type
        by_type = {}
        for doc_type in DocumentType:
            count = await db.documents.count_documents({"content_type": doc_type.value})
            by_type[doc_type.value] = count
        
        # By model
        pipeline = [
            {"$match": {"model_id": {"$exists": True, "$ne": None}}},
            {"$group": {"_id": "$model_id", "count": {"$sum": 1}}}
        ]
        model_counts = {}
        async for doc in db.documents.aggregate(pipeline):
            model_counts[doc["_id"]] = doc["count"]
        by_model = model_counts
        
        # Total chunks
        total_chunks = await db.knowledge_chunks.count_documents({})
        
        # Total words
        pipeline = [{"$group": {"_id": None, "total": {"$sum": "$word_count"}}}]
        word_result = await db.documents.aggregate(pipeline).to_list(1)
        total_words = word_result[0]["total"] if word_result else 0
        
        # Average extraction confidence
        pipeline = [
            {"$match": {"extraction_confidence": {"$exists": True}}},
            {"$group": {"_id": None, "avg": {"$avg": "$extraction_confidence"}}}
        ]
        conf_result = await db.documents.aggregate(pipeline).to_list(1)
        avg_confidence = conf_result[0]["avg"] if conf_result else None
        
        return DocumentStats(
            total_documents=total,
            by_status=by_status,
            by_type=by_type,
            by_model=by_model,
            total_chunks=total_chunks,
            total_words=total_words,
            avg_extraction_confidence=avg_confidence,
        )
    
    except Exception as e:
        logger.error(f"Failed to get stats: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to get stats: {str(e)}")


@router.get("/documents/{document_id}/download")
async def download_document(document_id: str):
    """Download the original file for a document."""
    try:
        from bson.objectid import ObjectId
        db = get_mongodb_client().db
        
        doc = await db.documents.find_one({"_id": ObjectId(document_id)})
        if not doc:
            raise HTTPException(status_code=404, detail="Document not found")
        
        file_path = doc.get("file_path")
        if not file_path or not os.path.exists(file_path):
            raise HTTPException(status_code=404, detail="File not found")
        
        return FileResponse(
            path=file_path,
            filename=doc.get("title", "download"),
            media_type=doc.get("mime_type", "application/octet-stream"),
        )
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to download document: {e}")
        raise HTTPException(status_code=500, detail=f"Download failed: {str(e)}")


@router.post("/documents/{document_id}/reprocess", response_model=ExtractionResult)
async def reprocess_document(document_id: str, request: ExtractionRequest):
    """Reprocess a document (re-extract and re-chunk)."""
    try:
        from bson.objectid import ObjectId
        db = get_mongodb_client().db
        
        # Get document
        doc = await db.documents.find_one({"_id": ObjectId(document_id)})
        if not doc:
            raise HTTPException(status_code=404, detail="Document not found")
        
        # Delete existing chunks from Qdrant and MongoDB
        try:
            qdrant = get_qdrant_client()
            collection = qdrant.get_collection()
            qdrant.client.delete(
                collection_name=collection,
                points_selector={"filter": {"must": [{"key": "document_id", "match": {"value": document_id}}]}}
            )
        except:
            pass
        
        await db.knowledge_chunks.delete_many({"document_id": document_id})
        
        # Reset status
        await db.documents.update_one(
            {"_id": ObjectId(document_id)},
            {
                "$set": {
                    "status": DocumentStatus.PENDING.value,
                    "chunk_count": 0,
                    "updated_at": datetime.utcnow(),
                }
            }
        )
        
        # Reprocess
        return await extract_document(document_id, request)
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to reprocess document: {e}")
        raise HTTPException(status_code=500, detail=f"Reprocess failed: {str(e)}")
